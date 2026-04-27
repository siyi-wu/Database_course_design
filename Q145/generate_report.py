from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt


OUTPUT_PATH = r"E:\Xiaomi Cloud\学习资料\数据库\大作业\实验室设备管理系统_个人报告_1_4_5.docx"


def set_document_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(12)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(18)


def add_subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(14 if level == 1 else 12)


def add_paragraph(doc: Document, text: str, first_line_indent: bool = True) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("• " + text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)


def add_code(doc: Document, code: str) -> None:
    for line in code.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
        run.font.size = Pt(10.5)


def add_table_from_rows(doc: Document, headers, rows) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)


def build_report() -> Document:
    doc = Document()
    set_document_style(doc)

    add_title(doc, "数据库大作业个人报告")
    add_subtitle(doc, "题目：实验室设备管理系统（本人负责内容：1、4、5）")
    add_subtitle(doc, "数据库：MySQL 8.0 / InnoDB")

    add_heading(doc, "一、本人负责内容")
    add_paragraph(doc, "根据小组分工，本人主要负责三部分内容：第一，建立至少五张表并保证表之间存在参照关系；第二，编写两个使用游标对数据库表进行操作的存储过程；第三，建立至少两个二级索引并完成相关查询性能分析。")

    add_heading(doc, "二、系统数据表与参照关系")
    add_paragraph(doc, "本系统以实验室设备管理为业务背景，围绕设备、类别、实验室、用户和借用记录五类核心对象建立数据库模式。最终共设计五张业务表，分别为 categories、labrooms、users、equipments 和 borrowrecords。")
    add_paragraph(doc, "各表之间的参照关系如下：设备表 equipments 通过 category_id 关联设备类别表 categories，通过 room_id 关联实验室表 labrooms；借用记录表 borrowrecords 通过 equip_id 关联设备表，通过 user_id 关联用户表。由此形成完整的数据联系网络。")

    add_table_from_rows(
        doc,
        ["表名", "主键", "主要字段", "说明"],
        [
            ["categories", "category_id", "category_name, description", "设备类别信息"],
            ["labrooms", "room_id", "room_name, location, admin_name", "实验室房间信息"],
            ["users", "user_id", "user_name, role, contact", "借用人信息"],
            ["equipments", "equip_id", "equip_name, category_id, room_id, status, price, purchase_date", "设备主数据"],
            ["borrowrecords", "record_id", "equip_id, user_id, borrow_date, plan_return_date, actual_return_date", "设备借还记录"],
        ],
    )

    add_paragraph(doc, "在新建的阿里云 MySQL 实例中已完成真实落库验证，外键关系包括：fk_equipments_category、fk_equipments_room、fk_borrowrecords_equipment 和 fk_borrowrecords_user。数据库能够保证借用记录、设备、用户、房间之间的一致性。")

    add_heading(doc, "三、两个游标存储过程设计")
    add_paragraph(doc, "为了满足课程对游标的要求，并体现数据库对逐行处理业务逻辑的支持，系统中设计了两个基于游标的存储过程。")

    add_heading(doc, "3.1 sp_sync_equipment_status()", level=2)
    add_paragraph(doc, "该过程用于批量同步设备状态。过程运行时，先逐条读取 equipments 表中的设备记录，再根据 borrowrecords 表判断当前设备是否存在尚未归还的借用记录。若设备处于 Maintenance 或 Scrapped 状态，则不进行修改；若设备存在未归还记录，则更新为 Borrowed；否则更新为 Available。")
    add_paragraph(doc, "该过程解决的是“设备状态字段可能与借用记录不一致”的问题，适用于系统批量校验或数据修复场景。实际测试中调用 CALL sp_sync_equipment_status(); 可以成功执行，执行耗时约为 0.363 秒。")

    add_heading(doc, "3.2 sp_batch_transfer_available_equipment(p_from_room_id, p_to_room_id, p_limit)", level=2)
    add_paragraph(doc, "该过程用于按条件批量调拨设备。过程先使用游标扫描指定实验室中状态为 Available 的设备，再按 purchase_date 和 equip_id 顺序逐台更新房间编号 room_id，直到达到给定上限 p_limit。")
    add_paragraph(doc, "该过程体现了游标在“逐行处理 + 过程控制”中的作用，适用于实验室之间的设备调拨业务。实际测试中调用 CALL sp_batch_transfer_available_equipment(1, 2, 5); 后，房间 1 的设备数量由 300 降为 295，房间 2 的设备数量由 300 升为 305，说明过程执行成功。")

    add_heading(doc, "四、索引设计")
    add_paragraph(doc, "为了提高高频查询效率，系统中建立了三个二级索引，其中前两个是报告重点分析对象。")
    add_bullet(doc, "idx_equipments_category_status(category_id, status)：建立在 equipments 表上，用于优化按设备类别和设备状态的联合查询。")
    add_bullet(doc, "idx_borrowrecords_user_borrow_date(user_id, borrow_date DESC)：建立在 borrowrecords 表上，用于优化按用户查询最近借用记录的排序型查询。")
    add_bullet(doc, "idx_borrowrecords_equip_active(equip_id, actual_return_date, plan_return_date)：建立在 borrowrecords 表上，用于优化按设备查询未归还记录、逾期记录等状态判断。")

    add_heading(doc, "五、索引性能分析")
    add_paragraph(doc, "为了更清楚地观察索引前后的差异，另外构建了两张基准测试表：benchmark_equipments 和 benchmark_borrowrecords。其中 benchmark_equipments 含 200000 行数据，benchmark_borrowrecords 含 1000000 行数据。测试时分别使用 IGNORE INDEX 强制走无索引方案，使用 FORCE INDEX 强制走目标索引方案。")

    add_heading(doc, "5.1 测试查询一：按类别和状态查询设备", level=2)
    add_code(
        doc,
        """
SELECT SQL_NO_CACHE equip_id, equip_name, status
FROM benchmark_equipments
WHERE category_id = 3
  AND status = 'Available';
        """,
    )
    add_paragraph(doc, "无索引时，执行计划为 Table scan on benchmark_equipments，即扫描整张设备表。EXPLAIN ANALYZE 显示扫描了 200000 行，平均测试耗时约为 473.908 ms。")
    add_paragraph(doc, "建立复合索引 idx_bm_equipments_category_status(category_id, status) 后，执行计划变为 Index lookup，平均测试耗时约为 349.980 ms。")
    add_paragraph(doc, "该查询的索引收益存在但不算特别夸张，主要原因是查询结果本身较多，共返回 30291 行。即使索引减少了定位成本，数据库仍然需要回表并向客户端返回大量结果，因此整体耗时还会受到结果传输影响。")

    add_heading(doc, "5.2 测试查询二：查询某用户最近借用记录", level=2)
    add_code(
        doc,
        """
SELECT SQL_NO_CACHE record_id, user_id, equip_id, borrow_date
FROM benchmark_borrowrecords
WHERE user_id = 42
ORDER BY borrow_date DESC
LIMIT 10;
        """,
    )
    add_paragraph(doc, "无索引时，执行计划为 Table scan + Sort，即先扫描 1000000 行借用记录，再按 borrow_date 倒序排序。平均测试耗时约为 166.785 ms。")
    add_paragraph(doc, "建立复合索引 idx_bm_borrowrecords_user_borrow_date(user_id, borrow_date DESC) 后，执行计划变为直接使用索引查找，不再进行全表扫描和显式排序。平均测试耗时降为 28.555 ms。")
    add_paragraph(doc, "该查询的优化效果明显，原因在于该索引同时支持了 where 条件、order by 排序以及 limit 取前 10 条三个需求，因此可以非常高效地返回结果。")

    add_heading(doc, "5.3 性能对比表", level=2)
    add_table_from_rows(
        doc,
        ["查询", "数据规模", "无索引平均耗时", "有索引平均耗时", "优化情况"],
        [
            ["设备类别+状态查询", "200000 行设备", "473.908 ms", "349.980 ms", "由全表扫描优化为复合索引查找"],
            ["用户最近借用记录查询", "1000000 行借用记录", "166.785 ms", "28.555 ms", "由全表扫描+排序优化为复合索引查找"],
        ],
    )

    add_heading(doc, "六、为什么部分查询提升不如预期")
    add_paragraph(doc, "索引并不是在所有场景下都能带来同等幅度的提升。对于“按类别和状态查询设备”这一测试，由于 category_id 和 status 的取值种类较少，条件选择性不高，同时查询需要返回 30291 行数据，因此即便已经使用索引，数据库仍需要处理大量结果。")
    add_paragraph(doc, "相比之下，“查询某用户最近借用记录”只需要返回少量结果，并且存在排序与 limit 条件，复合索引可以同时完成过滤、排序和裁剪，因此优化效果更显著。由此可见，索引提升是否明显，与数据规模、查询选择性、返回结果行数和排序需求都密切相关。")

    add_heading(doc, "七、进一步优化空间")
    add_paragraph(doc, "后续若需要进一步优化，还可以从以下方面入手：一是对展示类查询增加分页条件 limit，减少一次性返回的记录数量；二是根据更真实的业务查询模式继续设计更贴近业务的复合索引，例如 category_id、status、room_id 的联合索引；三是在查询字段固定时考虑覆盖索引，减少回表次数。")

    add_heading(doc, "八、结论")
    add_paragraph(doc, "本人完成了实验室设备管理系统中 5 张核心数据表及其参照关系设计，编写并验证了 2 个基于游标的存储过程，建立了 3 个二级索引，并通过大规模基准数据完成了索引前后性能分析。测试结果表明，合理建立复合索引能够显著改善数据库查询效率，尤其在带有筛选、排序和限制返回条数的查询场景中，优化效果更加明显。")

    add_heading(doc, "附录：关键 SQL 对象名称")
    add_bullet(doc, "存储过程：sp_sync_equipment_status")
    add_bullet(doc, "存储过程：sp_batch_transfer_available_equipment")
    add_bullet(doc, "正式业务索引：idx_equipments_category_status")
    add_bullet(doc, "正式业务索引：idx_borrowrecords_user_borrow_date")
    add_bullet(doc, "正式业务索引：idx_borrowrecords_equip_active")
    add_bullet(doc, "基准测试索引：idx_bm_equipments_category_status")
    add_bullet(doc, "基准测试索引：idx_bm_borrowrecords_user_borrow_date")

    return doc


if __name__ == "__main__":
    document = build_report()
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)
